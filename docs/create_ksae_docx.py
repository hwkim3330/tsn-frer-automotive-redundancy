#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_ksae_paper():
    # Create new document
    doc = docx.Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title
    title = doc.add_heading('TSN Switch의 FRER 기능을 이용한 자동차 아키텍처 이중화 기술 및 성능 평가', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add author info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run('저자: 김현우1)*\n소속: 한국전자기술연구원 모빌리티플랫폼연구센터(KETI Mobility Platform Research Center)1)\n교신저자 이메일: hwkim3330@keti.re.kr')
    author_run.font.size = Pt(12)

    # Add separator
    doc.add_paragraph('---')

    # Add English title
    eng_title = doc.add_heading('Evaluation of Automotive Architecture Redundancy and Performance Using TSN Switch FRER Function', level=2)
    eng_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add English author info
    eng_author_para = doc.add_paragraph()
    eng_author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eng_author_run = eng_author_para.add_run('Author: Hyunwoo Kim1)*\nAffiliation: Korea Electronics Technology Institute, Mobility Platform Research Center (KETI)1)\nCorresponding Author Email: hwkim3330@keti.re.kr')
    eng_author_run.font.size = Pt(12)

    # Add separator
    doc.add_paragraph('---')

    # Add Korean abstract
    abstract_heading = doc.add_heading('초록(Abstract)', level=2)
    
    # Abstract text in multiple paragraphs for better formatting
    abstract_paragraphs = [
        '현대 자동차는 자율주행 시스템, ADAS(Advanced Driver Assistance System), 전동화 제어 시스템 등 안전 크리티컬 애플리케이션의 급속한 증가로 인해 네트워크 통신의 신뢰성과 실시간성에 대한 요구가 급격히 높아지고 있다. 특히 브레이크 바이 와이어, 조향 제어, 센서 융합과 같은 안전 필수 기능에서는 네트워크 단절이나 링크 오류가 발생하더라도 서비스의 연속성을 보장해야 하며, 단일 장애점(Single Point of Failure)으로 인한 시스템 전체 마비를 방지해야 한다. 그러나 기존의 단일 경로 기반 이더넷 통신은 링크 장애나 스위치 오류 시 패킷 손실과 서비스 중단이 직접적으로 발생하여, ISO 26262와 같은 자동차 안전 표준이 요구하는 높은 신뢰성 수준을 충족하기 어려운 한계가 있다.',
        
        '이러한 문제를 근본적으로 해결하기 위한 핵심 기술이 IEEE 802.1CB에서 정의한 FRER(Frame Replication and Elimination for Reliability)이다. FRER은 송신단에서 중요한 프레임을 여러 독립적인 경로로 복제하여 전송하고, 수신단에서 중복된 프레임을 지능적으로 제거하여 원본 스트림을 복원함으로써 네트워크 장애에 대한 투명한 복구 메커니즘을 제공한다. 본 연구는 실제 TSN(Time-Sensitive Networking) 스위치 환경에서 FRER의 전체 동작 과정—스트림 식별, 시퀀스 생성, 프레임 복제, 중복 제거—을 구현하고, 와이어샤크(Wireshark) 패킷 분석을 통해 그 동작을 상세히 검증하였다.',
        
        '실험 환경에서는 Microchip LAN9662 TSN 스위치와 Raspberry Pi CM4를 기반으로 하는 테스트베드를 구축하였다. 송신 진입 스위치에서는 VCAP(Versatile Content-Aware Processor) 기반 규칙을 사용하여 UDP 트래픽을 선별적으로 식별하고, 이를 고유한 스트림 식별자(ISDX)로 분류하였다. 스위치 하드웨어는 식별된 스트림에 대해 자동으로 R-TAG(Redundancy TAG, EtherType 0xF1C1)를 삽입하고 시퀀스 번호를 순차적으로 부여한 뒤, 동일한 프레임을 두 개의 물리적으로 독립된 경로(eth2, eth3)로 복제하여 전송한다. 각 복제 프레임은 동일한 시퀀스 번호를 가지며, 이는 수신단에서 중복 검출의 핵심 기준으로 사용된다.',
        
        '수신측에서는 각 경로에서 도착하는 프레임들을 개별 멤버 스트림(Member Stream)으로 처리한 뒤, 이들을 단일 컴파운드 스트림(Compound Stream)으로 병합하여 시퀀스 복구(Sequence Recovery) 기능을 수행한다. 동일한 시퀀스 번호를 가진 프레임 중 먼저 도착한 것만을 유효한 프레임으로 채택하고, 지연되어 나중에 도착한 중복 프레임은 자동으로 폐기한다. 이 과정은 상위 애플리케이션에게 완전히 투명하게 이루어지므로, 엔드 노드는 마치 단일 경로에서 안정적으로 데이터를 수신하는 것처럼 인식하게 된다.',
        
        '실제 트래픽 생성은 mausezahn 도구를 사용하여 다양한 주기의 UDP 스트림을 생성하였으며, 10.0.100.1에서 10.0.100.2 및 10.0.100.3으로의 연속적인 패킷 전송을 통해 이중화 시나리오를 시뮬레이션하였다. 와이어샤크를 통한 패킷 캡처 분석에서는 Ethernet II, R-TAG, IPv4, UDP 형태의 계층 구조가 정확히 관찰되었으며, R-TAG 내부의 시퀀스 번호가 18663에서 18676으로 단조 증가하는 패턴을 확인할 수 있었다. 특히 동일한 페이로드가 서로 다른 목적지로 전송될 때에도 각각 고유한 시퀀스 번호가 부여되어 독립적으로 처리되는 것을 확인하였다.',
        
        '링크 장애 시뮬레이션 실험에서는 Primary 경로를 의도적으로 차단했을 때 Secondary 경로를 통한 무중단 서비스 연속성을 확인하였으며, 복구 과정에서 패킷 손실이나 순서 뒤바뀜 없이 정상적인 데이터 전달이 유지됨을 입증하였다. 또한 두 경로가 모두 활성화된 상태에서는 수신측 스위치가 중복 프레임을 효과적으로 제거하여 단일 스트림으로 복원하는 과정이 올바르게 동작함을 확인하였다.',
        
        '본 연구의 결과는 FRER 기술이 자동차 이더넷 환경에서 네트워크 신뢰성 향상을 위한 효과적이고 실용적인 솔루션임을 실증적으로 보여준다. 특히 하드웨어 기반의 프레임 복제와 시퀀스 관리가 소프트웨어 오버헤드 없이 수행되므로, 실시간 요구사항이 엄격한 자동차 시스템에 적합하다. 또한 상위 애플리케이션의 수정 없이 네트워크 계층에서 투명하게 이중화를 제공함으로써, 기존 시스템과의 호환성을 유지하면서도 안전성을 대폭 향상시킬 수 있다는 장점을 확인하였다. 향후 연구에서는 대규모 멀티홉 네트워크에서의 확장성과 다양한 트래픽 클래스 간의 상호작용에 대한 추가적인 분석이 필요할 것으로 판단된다.'
    ]
    
    for para_text in abstract_paragraphs:
        para = doc.add_paragraph(para_text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Add spacing between paragraphs
        para.paragraph_format.space_after = Pt(6)

    # Add separator
    doc.add_paragraph('---')

    # Add keywords
    keywords_heading = doc.add_heading('키워드 (Key words)', level=2)
    keywords_para = doc.add_paragraph()
    keywords_para.add_run('한글: TSN, FRER, 이중화, 자동차 이더넷, 신뢰성, 자율주행\n')
    keywords_para.add_run('영문: TSN, FRER, Redundancy, Automotive Ethernet, Reliability, Autonomous Driving')

    # Add separator
    doc.add_paragraph('---')

    # Add experimental setup section
    setup_heading = doc.add_heading('실험 설정 및 구현 상세 (부록)', level=2)
    
    # Traffic generation section
    traffic_heading = doc.add_heading('1. 트래픽 생성 (송신 노드)', level=3)
    traffic_code = '''# UDP 스트림 생성 - 목적지 10.0.100.2
sudo mausezahn eth0 -c 0 -d 1000000 \\
  -t udp sp=5000,dp=5000 -A 10.0.100.1 -B 10.0.100.2

# UDP 스트림 생성 - 목적지 10.0.100.3  
sudo mausezahn eth0 -c 0 -d 1000000 \\
  -t udp sp=5000,dp=5000 -A 10.0.100.1 -B 10.0.100.3

# 옵션 설명:
# -c 0: 무한 전송
# -d: 패킷 간격 (나노초 단위)
# -t udp: UDP 프로토콜 사용
# sp/dp: 소스/목적지 포트 번호'''
    
    traffic_para = doc.add_paragraph(traffic_code)
    traffic_para.style = 'Normal'

    # VCAP section
    vcap_heading = doc.add_heading('2. VCAP 기반 스트림 식별', level=3)
    vcap_code = '''# UDP/IPv4 트래픽을 ISDX=1로 분류
vcap add 1 is1 10 0 VCAP_KFS_5TUPLE_IP4 \\
  IF_IGR_PORT_MASK 0x001 0x1ff L3_IP_PROTO 17 0xff \\
  VCAP_AFS_S1 ISDX_REPLACE_ENA 1 ISDX_ADD_VAL 1

# 설정 확인
vcap get 1
# 결과: Keyset: VCAP_KFS_5TUPLE_IP4, Action: ISDX=1'''
    
    vcap_para = doc.add_paragraph(vcap_code)
    vcap_para.style = 'Normal'

    # FRER flow section
    frer_heading = doc.add_heading('3. FRER 복제 플로우 설정', level=3)
    frer_code = '''# ISDX=1 스트림에 시퀀스 생성 및 이중 경로 분기 활성화
frer iflow 1 --generation 1 --dev1 eth2 --dev2 eth3

# 설정 상태 확인
frer iflow 1
# 결과:
# ms_enable: 0
# generation: 1 (시퀀스 생성 활성화)
# dev1: eth2 (첫 번째 출력 포트)
# dev2: eth3 (두 번째 출력 포트)'''
    
    frer_para = doc.add_paragraph(frer_code)
    frer_para.style = 'Normal'

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('\n본 문서는 2025 KSAE 춘계학술대회 논문 초록입니다.')
    footer_run.font.size = Pt(10)
    footer_run.bold = True

    # Save document
    doc.save('KSAE_2025_FRER_논문초록.docx')
    print('DOCX 파일이 성공적으로 생성되었습니다: KSAE_2025_FRER_논문초록.docx')

if __name__ == '__main__':
    create_ksae_paper()